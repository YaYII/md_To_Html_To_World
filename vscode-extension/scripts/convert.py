#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Markdown 转 Word 文档转换脚本
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Dict, Any

try:
    import markdown
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENTATION
    from bs4 import BeautifulSoup
except ImportError as e:
    print(f"缺少必要的依赖: {e}", file=sys.stderr)
    sys.exit(1)

class MarkdownToWordConverter:
    """Markdown 转 Word 转换器"""
    
    def __init__(self, config: Dict[str, Any]):
        """
        初始化转换器
        
        Args:
            config: 转换配置
        """
        self.config = config
        self.document = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档格式"""
        # 设置页面大小和方向
        section = self.document.sections[0]
        if self.config['format']['orientation'] == 'landscape':
            section.orientation = WD_ORIENTATION.LANDSCAPE
        
        # 设置页边距
        margins = self.config['format']['margins']
        section.top_margin = Inches(margins['top'])
        section.bottom_margin = Inches(margins['bottom'])
        section.left_margin = Inches(margins['left'])
        section.right_margin = Inches(margins['right'])
        
        # 设置默认字体
        style = self.document.styles['Normal']
        font = style.font
        font.name = self.config['style']['font_family']
        font.size = Pt(self.config['style']['font_size'])
    
    def convert(self, input_file: str, output_file: str):
        """
        执行转换
        
        Args:
            input_file: 输入文件路径
            output_file: 输出文件路径
        """
        try:
            # 读取 Markdown 文件
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # 转换为 HTML
            html = markdown.markdown(
                markdown_text,
                extensions=[
                    'markdown.extensions.extra',
                    'markdown.extensions.toc',
                    'markdown.extensions.tables',
                    'markdown.extensions.fenced_code'
                ]
            )
            
            # 解析 HTML
            soup = BeautifulSoup(html, 'html.parser')
            
            # 添加目录
            if self.config['toc']['include']:
                self._add_toc()
            
            # 转换内容
            self._convert_content(soup)
            
            # 保存文档
            self.document.save(output_file)
            
        except Exception as e:
            print(f"转换过程出错: {e}", file=sys.stderr)
            sys.exit(1)
    
    def _add_toc(self):
        """添加目录"""
        # 添加目录标题
        title = self.document.add_heading(
            self.config['toc']['title'],
            level=1
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加目录域代码
        para = self.document.add_paragraph()
        run = para.add_run()
        run.add_field('TOC', f'\\o "1-{self.config["toc"]["depth"]}"')
        
        # 添加分页符
        self.document.add_page_break()
    
    def _convert_content(self, soup: BeautifulSoup):
        """
        转换文档内容
        
        Args:
            soup: BeautifulSoup 对象
        """
        for element in soup.children:
            if element.name is None:
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                self.document.add_heading(element.get_text(), level=level)
                
            elif element.name == 'p':
                para = self.document.add_paragraph()
                for child in element.children:
                    if child.name == 'img' and self.config['advanced']['preserve_images']:
                        self._add_image(para, child)
                    elif child.name == 'a' and self.config['advanced']['preserve_links']:
                        self._add_link(para, child)
                    else:
                        para.add_run(child.get_text())
                        
            elif element.name == 'ul':
                self._add_list(element, is_numbered=False)
                
            elif element.name == 'ol':
                self._add_list(element, is_numbered=True)
                
            elif element.name == 'table':
                self._add_table(element)
    
    def _add_image(self, paragraph, img):
        """添加图片"""
        src = img.get('src', '')
        if not src:
            return
            
        try:
            # 处理相对路径
            if not src.startswith(('http://', 'https://')):
                src = os.path.join(os.path.dirname(input_file), src)
            
            # 添加图片
            width = self.config['advanced'].get('image_max_width', 800)
            paragraph.add_run().add_picture(src, width=Pt(width))
            
        except Exception as e:
            print(f"添加图片失败 ({src}): {e}", file=sys.stderr)
    
    def _add_link(self, paragraph, link):
        """添加链接"""
        text = link.get_text()
        url = link.get('href', '')
        if url:
            paragraph.add_hyperlink(url, text)
        else:
            paragraph.add_run(text)
    
    def _add_list(self, list_elem, is_numbered=False):
        """添加列表"""
        for item in list_elem.find_all('li', recursive=False):
            para = self.document.add_paragraph(
                style='List Number' if is_numbered else 'List Bullet'
            )
            para.add_run(item.get_text())
    
    def _add_table(self, table):
        """添加表格"""
        rows = table.find_all('tr')
        if not rows:
            return
            
        # 获取列数
        cols = len(rows[0].find_all(['td', 'th']))
        if cols == 0:
            return
            
        # 创建表格
        table = self.document.add_table(rows=0, cols=cols)
        table.style = 'Table Grid'
        
        # 添加行
        for row in rows:
            cells = row.find_all(['td', 'th'])
            if len(cells) != cols:
                continue
                
            row_cells = table.add_row().cells
            for i, cell in enumerate(cells):
                row_cells[i].text = cell.get_text().strip()

def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='Markdown 转 Word 文档转换工具')
    parser.add_argument('--input', required=True, help='输入 Markdown 文件路径')
    parser.add_argument('--output', required=True, help='输出 Word 文件路径')
    parser.add_argument('--config', required=True, help='配置文件路径')
    
    args = parser.parse_args()
    
    try:
        # 读取配置
        with open(args.config, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        # 创建转换器
        converter = MarkdownToWordConverter(config)
        
        # 执行转换
        converter.convert(args.input, args.output)
        
    except Exception as e:
        print(f"转换失败: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main() 