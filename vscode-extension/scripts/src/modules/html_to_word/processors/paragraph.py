"""
段落处理器模块
处理HTML段落相关元素
"""

import re
from typing import Dict, Any, Optional, List, Union
from bs4 import Tag
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .base import BaseProcessor
from .inline import InlineProcessor

class ParagraphProcessor(BaseProcessor):
    """
    /**
     * 段落处理器
     * 
     * 处理HTML段落相关元素，如p、blockquote等
     */
    """
    
    def can_process(self, element: Tag) -> bool:
        """
        /**
         * 检查是否可以处理指定元素
         * 
         * @param {Tag} element - HTML元素
         * @returns {bool} 是否可以处理该元素
         */
        """
        paragraph_tags = ['p', 'blockquote', 'div', 'section', 'article', 'main', 'header', 'footer']
        return element.name in paragraph_tags
    
    def process(self, element: Tag) -> Union[Paragraph, List[Paragraph], None]:
        """
        /**
         * 处理段落元素
         * 
         * @param {Tag} element - 段落HTML元素
         * @returns {Union[Paragraph, List[Paragraph], None]} 处理结果
         */
        """
        if self.debug_mode:
            self.logger.debug(f"处理段落元素: <{element.name}> {element.get('id', '')} {element.get('class', '')}")
            
        if element.name == 'blockquote':
            return self._process_blockquote(element)
        else:
            return self._process_paragraph(element)
    
    def _process_paragraph(self, element: Tag) -> Optional[Paragraph]:
        """
        /**
         * 处理段落元素
         * 
         * @param {Tag} element - 段落HTML元素
         * @returns {Optional[Paragraph]} 处理后的段落对象
         */
        """
        # 检查是否使用br标签模拟列表
        if self._handle_br_formatted_list(element):
            # 如果成功处理为列表样式，直接返回
            if self.debug_mode:
                self.logger.debug(f"段落被检测为使用br标签模拟的列表，已特殊处理")
            return None
            
        # 创建段落并应用样式
        p = self.document.add_paragraph()
        self.style_manager.apply_paragraph_format(p)
        
        # 处理对齐方式
        if 'align' in element.attrs:
            self.style_manager.apply_paragraph_format(p, element['align'].lower())
            if self.debug_mode:
                self.logger.debug(f"应用段落对齐方式: {element['align'].lower()}")
        
        # 使用内联元素处理器处理内容
        inline_processor = InlineProcessor(self.document, self.style_manager)
        inline_processor.process_inline_elements(element, p)
        
        if self.debug_mode:
            self.logger.debug(f"段落处理完成，文本内容: {p.text[:30]}{'...' if len(p.text) > 30 else ''}")
        return p
    
    def _process_blockquote(self, element: Tag) -> List[Paragraph]:
        """
        /**
         * 处理引用块元素
         * 
         * @param {Tag} element - 引用块HTML元素
         * @returns {List[Paragraph]} 处理后的段落对象列表
         */
        """
        if self.debug_mode:
            self.logger.debug(f"处理引用块元素: <{element.name}>")
            
        paragraphs = []
        
        # 为引用块中的每个子元素创建段落
        for child in element.children:
            if isinstance(child, Tag):
                if child.name == 'p':
                    # 创建带引用样式的段落
                    p = self.document.add_paragraph()
                    self.style_manager.apply_quote_format(p)
                    
                    # 添加引用样式的竖线
                    p.add_run('│ ')
                    
                    # 处理段落内容
                    inline_processor = InlineProcessor(self.document, self.style_manager)
                    inline_processor.process_inline_elements(child, p)
                    
                    paragraphs.append(p)
                    if self.debug_mode:
                        self.logger.debug(f"处理引用块中的段落: {p.text[:30]}{'...' if len(p.text) > 30 else ''}")
                else:
                    # 处理其他类型的子元素
                    if self.debug_mode:
                        self.logger.debug(f"处理引用块中的其他元素: <{child.name}>")
                    # 获取处理器
                    from ..element_factory import ElementProcessorFactory
                    factory = ElementProcessorFactory(self.document, self.style_manager)
                    processor = factory.get_processor(child)
                    
                    if processor:
                        result = processor.process(child)
                        if isinstance(result, list):
                            paragraphs.extend(result)
                        elif result:
                            paragraphs.append(result)
        
        if self.debug_mode:
            self.logger.debug(f"引用块处理完成，生成了 {len(paragraphs)} 个段落")
        return paragraphs
    
    def _handle_br_formatted_list(self, element: Tag) -> bool:
        """
        /**
         * 处理使用br标签模拟的列表
         * 
         * @param {Tag} element - HTML元素
         * @returns {bool} 是否成功处理为列表
         */
        """
        # 查找所有换行标签
        br_tags = element.find_all('br')
        
        # 如果没有换行标签，直接返回False
        if not br_tags:
            return False
            
        # 检查是否可能是使用br标签模拟的列表
        # 一些常见的列表项前缀模式
        list_prefixes = [
            re.compile(r'^\s*(\d+)[\.\)]\s'),  # 数字列表: "1. " 或 "1) "
            re.compile(r'^\s*[•\-\*]\s'),      # 无序列表: "• " 或 "- " 或 "* "
            re.compile(r'^\s*[a-zA-Z][\.\)]\s')  # 字母列表: "a. " 或 "A) "
        ]
        
        # 拆分文本，检查每行是否符合列表项格式
        lines = []
        current_line = ""
        
        # 递归收集所有文本，保持顺序
        def collect_text(node):
            nonlocal current_line
            if isinstance(node, Tag):
                if node.name == 'br':
                    lines.append(current_line)
                    current_line = ""
                else:
                    for child in node.children:
                        collect_text(child)
            else:
                current_line += str(node)
        
        # 收集第一级子元素的文本
        for child in element.children:
            collect_text(child)
        
        # 添加最后一行
        if current_line:
            lines.append(current_line)
        
        # 检查是否有行匹配列表项模式
        list_like = False
        for line in lines:
            for pattern in list_prefixes:
                if pattern.search(line):
                    list_like = True
                    break
            if list_like:
                break
                
        # 如果像列表，特殊处理每一行
        if list_like and lines:
            for line in lines:
                if line.strip():
                    # 为每个"列表项"创建一个新的段落
                    p = self.document.add_paragraph()
                    p.paragraph_format.line_spacing = self.style_manager.line_spacing
                    p.paragraph_format.space_after = self.style_manager.paragraph_spacing / 2
                    
                    # 添加文本
                    run = p.add_run(line.strip())
                    self.style_manager.apply_default_style(run)
            return True
            
        # 如果不是类似列表，返回False
        return False 