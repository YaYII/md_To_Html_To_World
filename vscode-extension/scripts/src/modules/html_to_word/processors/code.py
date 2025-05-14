"""
代码块处理器模块
处理HTML代码块元素
"""

from typing import Dict, Any, Optional, List, Union
from bs4 import Tag
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from .base import BaseProcessor

class CodeProcessor(BaseProcessor):
    """
    /**
     * 代码块处理器
     * 
     * 处理HTML代码块元素，如pre和code
     */
    """
    
    def can_process(self, element: Tag) -> bool:
        """
        /**
         * 检查是否可以处理指定元素
         * 
         * @param {Tag} element - HTML元素
         * @returns {bool} 是否可以处理该元素
         */
        """
        code_tags = ['pre', 'code']
        return element.name in code_tags
    
    def process(self, element: Tag) -> Union[Paragraph, None]:
        """
        /**
         * 处理代码块元素
         * 
         * @param {Tag} element - 代码块HTML元素
         * @returns {Union[Paragraph, None]} 处理后的段落对象或None
         */
        """
        if self.debug_mode:
            self.logger.debug(f"处理代码块元素: <{element.name}>")
            
        # 查找代码元素，可能在pre内部嵌套了code标签
        code_element = element.find('code') or element
        code_text = code_element.get_text()
        
        # 确保代码块有内容
        if not code_text.strip():
            if self.debug_mode:
                self.logger.debug("代码块内容为空，忽略处理")
            return None
            
        # 处理代码文本，移除多余的空行并分割为行
        code_lines = code_text.strip().split('\n')
        
        if self.debug_mode:
            self.logger.debug(f"代码块包含 {len(code_lines)} 行代码")
        
        # 创建一个单列表格来容纳代码
        table = self.document.add_table(rows=1, cols=1)
        
        # 获取单元格
        cell = table.cell(0, 0)
        
        # 为单元格设置黑色背景
        # 使用XML直接设置黑色背景
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 移除现有的单元格阴影设置（如果有）
        existing_shading = tcPr.find(qn('w:shd'))
        if existing_shading is not None:
            tcPr.remove(existing_shading)
        
        # 创建并添加新的阴影元素 (黑色背景)
        bg_color = "1A1A1A"  # 深灰黑色
        shading_element = parse_xml(
            f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg_color}"/>'
        )
        tcPr.append(shading_element)
        
        # 设置表格边框为无边框
        borders_xml = f'<w:tcBorders {nsdecls("w")}>'
        for direction in ['top', 'left', 'bottom', 'right']:
            borders_xml += f'<w:{direction} w:val="single" w:sz="2" w:space="0" w:color="{bg_color}"/>'
        borders_xml += '</w:tcBorders>'
        
        # 移除现有的边框设置
        for old in tcPr.findall('.//{*}tcBorders'):
            tcPr.remove(old)
        
        # 添加新的边框设置
        tcPr.append(parse_xml(borders_xml))
        
        # 设置单元格内边距
        cell_margin = 6  # 6磅内边距
        margins_xml = f'<w:tcMar {nsdecls("w")}>'
        for direction in ['top', 'left', 'bottom', 'right']:
            margins_xml += f'<w:{direction} w:w="{cell_margin*20}" w:type="dxa"/>'
        margins_xml += '</w:tcMar>'
        
        # 移除现有的边距设置
        for old in tcPr.findall('.//{*}tcMar'):
            tcPr.remove(old)
        
        # 添加新的边距设置
        tcPr.append(parse_xml(margins_xml))
        
        # 清空单元格内容并添加代码文本
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        
        # 添加所有代码行，每行之间添加换行符
        for i, line in enumerate(code_lines):
            if i > 0:
                paragraph.add_run("\n")
            
            run = paragraph.add_run(line)
            # 使用等宽字体显示代码
            run.font.name = self.style_manager.code_font
            run.font.size = Pt(self.style_manager.code_size)
            # 设置为白色文字
            text_color = "FFFFFF"  # 纯白色
            run.font.color.rgb = RGBColor(255, 255, 255)
        
        if self.debug_mode:
            self.logger.debug(f"代码块处理完成，已创建黑底白字代码块")
            
        # 在表格后添加一个空段落，增加间距
        spacer = self.document.add_paragraph()
        return spacer