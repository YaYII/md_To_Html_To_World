"""
图片处理模块
处理HTML中的图片元素转换为Word文档
"""

import os
import tempfile
import requests
from typing import Optional, Dict, Any
from bs4 import Tag
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

from .base import BaseProcessor


class ImageProcessor(BaseProcessor):
    """
    /**
     * 图片处理器
     * 
     * 负责将HTML图片元素转换为Word文档中的图片
     */
    """
    
    def __init__(self, document: Document, style_manager):
        """
        /**
         * 初始化图片处理器
         * 
         * @param {Document} document - Word文档对象
         * @param {object} style_manager - 样式管理器对象
         */
        """
        super().__init__(document, style_manager)
        self.image_config = self.style_manager.config.get('images', {})
        self.default_image_width = self.image_config.get('default_width', 10)  # 单位：厘米
        self.default_image_height = self.image_config.get('default_height', None)  # 自动计算高度
        self.max_image_width = self.image_config.get('max_width', 15)  # 单位：厘米
        self.temp_dir = None
    
    def process(self, element: Tag) -> Optional[Paragraph]:
        """
        /**
         * 处理图片元素，将HTML图片转换为Word图片
         * 
         * @param {Tag} element - HTML图片标签元素
         * @returns {Optional[Paragraph]} 包含图片的段落对象或None
         */
        """
        try:
            # 获取图片源
            src = element.get('src', '')
            if not src:
                return None
                
            # 获取图片alt文本
            alt = element.get('alt', '')
            
            # 创建临时目录（如果尚未创建）
            if self.temp_dir is None:
                self.temp_dir = tempfile.mkdtemp()
                
            # 下载或读取图片文件
            img_path = self._get_image_path(src)
            if not img_path:
                return None
                
            # 获取图片尺寸设置
            width_attr = element.get('width', '')
            width = self._parse_size_attribute(width_attr, self.default_image_width)
            
            # 检查图片宽度是否超过最大值
            if width > self.max_image_width:
                width = self.max_image_width
                
            # 创建段落并插入图片
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 插入图片
            run = p.add_run()
            run.add_picture(img_path, width=Cm(width))
            
            # 如果有alt文本，添加图片说明
            if alt and self.image_config.get('show_caption', True):
                caption_p = self.document.add_paragraph(f"图 {alt}")
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.paragraph_format.space_after = Pt(12)
                
            return p
            
        except Exception as e:
            # 图片处理失败时返回None
            print(f"图片处理失败: {str(e)}")
            return None
    
    def _get_image_path(self, src: str) -> Optional[str]:
        """
        /**
         * 获取图片路径，支持本地文件和远程URL
         * 
         * @param {str} src - 图片源路径
         * @returns {Optional[str]} 图片的本地路径或None
         */
        """
        # 检查是否为本地文件路径
        if os.path.exists(src):
            return src
            
        # 检查是否为URL
        if src.startswith(('http://', 'https://')):
            try:
                # 从URL下载图片
                response = requests.get(src, stream=True, timeout=10)
                if response.status_code == 200:
                    # 提取文件名
                    filename = os.path.basename(src)
                    if not filename:
                        filename = f"image_{hash(src)}.jpg"
                        
                    # 保存到临时文件
                    img_path = os.path.join(self.temp_dir, filename)
                    with open(img_path, 'wb') as f:
                        for chunk in response.iter_content(1024):
                            f.write(chunk)
                    return img_path
            except Exception:
                pass
                
        # 如果是相对路径，尝试在指定目录中查找
        image_dirs = self.image_config.get('search_dirs', [])
        for dir_path in image_dirs:
            full_path = os.path.join(dir_path, src)
            if os.path.exists(full_path):
                return full_path
                
        return None
    
    def _parse_size_attribute(self, size_attr: str, default_value: float) -> float:
        """
        /**
         * 解析HTML大小属性
         * 
         * @param {str} size_attr - HTML大小属性字符串
         * @param {float} default_value - 默认值
         * @returns {float} 解析后的尺寸值（厘米）
         */
        """
        if not size_attr:
            return default_value
            
        try:
            # 如果是纯数字，假设为像素值
            if size_attr.isdigit():
                # 像素转厘米（近似转换）
                return float(size_attr) / 40  # 约40px = 1cm
                
            # 尝试解析带单位的值
            if size_attr.endswith('px'):
                return float(size_attr[:-2]) / 40
            elif size_attr.endswith('cm'):
                return float(size_attr[:-2])
            elif size_attr.endswith('in'):
                return float(size_attr[:-2]) * 2.54  # 英寸转厘米
            elif size_attr.endswith('%'):
                # 百分比值，基于最大宽度
                percent = float(size_attr[:-1]) / 100
                return self.max_image_width * percent
        except ValueError:
            pass
            
        return default_value
    
    def cleanup(self):
        """
        /**
         * 清理临时资源
         */
        """
        # 清理临时文件夹
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                import shutil
                shutil.rmtree(self.temp_dir)
                self.temp_dir = None
            except Exception as e:
                self.logger.error(f"清理临时文件失败: {e}")
                pass 