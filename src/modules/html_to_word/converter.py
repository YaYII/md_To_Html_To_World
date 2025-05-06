"""
HTML到Word转换器
提供HTML内容转换为Word文档的功能
"""

import os
import re
import time
import logging
from typing import Dict, Any, Optional, List, Union
import codecs
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .document_style import DocumentStyleManager
from .element_factory import ElementProcessorFactory
from ..html_elements_processor import HtmlElementsProcessor

class HtmlToWordConverter:
    """
    /**
     * HTML到Word转换器
     * 
     * 负责将HTML内容转换为格式化的Word文档
     */
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        /**
         * 初始化HTML到Word转换器
         * 
         * @param {Dict[str, Any]} config - 配置参数字典
         */
        """
        self.config = config
        self.document = None
        self.style_manager = DocumentStyleManager(config)
        self.processor_factory = None
        self.elements_processor = HtmlElementsProcessor(config)
        
        # 配置日志
        self.debug_mode = config.get('debug', {}).get('enabled', False)
        log_level = logging.DEBUG if self.debug_mode else logging.INFO
        
        # 创建logger
        self.logger = logging.getLogger('HtmlToWordConverter')
        self.logger.setLevel(log_level)
        
        # 如果没有处理器，添加一个控制台处理器
        if not self.logger.handlers:
            console_handler = logging.StreamHandler()
            console_handler.setLevel(log_level)
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(formatter)
            self.logger.addHandler(console_handler)
        
        self.logger.info("HTML到Word转换器初始化完成")
        if self.debug_mode:
            self.logger.debug(f"调试模式已启用，配置: {config}")
    
    def convert_file(self, input_file: str, output_file: str) -> Document:
        """
        /**
         * 将HTML文件转换为Word文档
         * 
         * @param {str} input_file - 输入HTML文件路径
         * @param {str} output_file - 输出Word文件路径
         * @returns {Document} 生成的Word文档对象
         */
        """
        self.logger.info(f"开始转换文件: {input_file} -> {output_file}")
        start_time = time.time()
        
        if not os.path.exists(input_file):
            self.logger.error(f"输入文件不存在: {input_file}")
            raise FileNotFoundError(f"输入文件不存在: {input_file}")
        
        try:
            with codecs.open(input_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
                self.logger.debug(f"成功读取HTML文件，大小: {len(html_content)} 字节")
                
            doc = self.convert_html(html_content)
            doc.save(output_file)
            
            elapsed_time = time.time() - start_time
            self.logger.info(f"文件转换完成，耗时: {elapsed_time:.2f} 秒")
            return doc
            
        except Exception as e:
            self.logger.error(f"转换过程中发生错误: {str(e)}", exc_info=True)
            raise
    
    def convert_html(self, html_content: str) -> Document:
        """
        /**
         * 将HTML内容转换为Word文档
         * 
         * @param {str} html_content - HTML格式的内容
         * @returns {Document} 生成的Word文档对象
         */
        """
        self.logger.info("开始转换HTML内容到Word")
        
        # 创建新文档
        self.document = Document()
        self.logger.debug("创建新文档对象")
        
        # 应用文档样式
        self.document = self.style_manager.setup_document(self.document)
        self.logger.debug("应用文档样式设置完成")
        
        # 检查是否需要生成目录
        if self.config.get('document', {}).get('generate_toc', False):
            self.logger.info("添加文档目录")
            self._add_table_of_contents()
        
        # 初始化处理器工厂
        self.processor_factory = ElementProcessorFactory(self.document, self.style_manager)
        self.logger.debug("初始化元素处理器工厂")
        
        # 解析HTML
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            body = soup.body or soup
            self.logger.debug(f"HTML解析完成，找到 {len(list(body.descendants))} 个元素")
        except Exception as e:
            self.logger.error(f"HTML解析失败: {str(e)}")
            raise
        
        # 处理主体内容
        self._process_body(body)
        
        self.logger.info("HTML内容转换完成")
        return self.document
    
    def _add_table_of_contents(self):
        """
        /**
         * 向文档添加目录
         * 使用Word字段代码生成目录，用户需要在Word中右键更新目录
         */
        """
        self.logger.debug("添加目录")
        
        # 添加"目录"标题段落
        toc_title = self.document.add_paragraph("目录")
        toc_title.alignment = 1  # 居中对齐
        toc_title.style = self.document.styles['Heading 1']
        
        # 添加空段落以插入目录字段
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        
        # 创建目录字段开始标记
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        # 创建目录字段指令
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        # 设置目录指令：包含1-3级标题，添加超链接，显示页码
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        # 创建目录字段分隔符
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # 创建提示文本
        t_element = OxmlElement('w:t')
        t_element.text = "右键点击此处更新目录"
        fldChar2.append(t_element)
        
        # 添加自动更新字段的标记
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        
        # 创建目录字段结束标记
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        # 将所有元素添加到run中
        r_element = run._r  # 获取底层XML元素
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(updateFields)
        r_element.append(fldChar4)
        
        # 添加分页符
        self.document.add_page_break()
        self.logger.info("目录添加完成")
    
    def _process_body(self, body: Tag):
        """
        /**
         * 处理HTML文档主体
         * 
         * @param {Tag} body - HTML文档主体元素
         */
        """
        self.logger.debug("开始处理文档主体")
        
        # 统计处理的元素
        processed_count = 0
        
        # 遍历所有直接子元素
        for child in body.children:
            if isinstance(child, Tag):
                self._process_element(child)
                processed_count += 1
                
        self.logger.debug(f"文档主体处理完成，共处理 {processed_count} 个顶级元素")
    
    def _process_element(self, element: Tag):
        """
        /**
         * 处理HTML元素
         * 
         * @param {Tag} element - HTML元素
         */
        """
        if self.debug_mode:
            self.logger.debug(f"处理元素: <{element.name}> {element.get('id', '')} {element.get('class', '')}")
        
        # 获取元素处理器
        try:
            processor = self.processor_factory.get_processor(element)
            
            if processor:
                # 使用处理器处理元素
                processor.process(element)
                if self.debug_mode:
                    self.logger.debug(f"使用 {processor.__class__.__name__} 处理元素 <{element.name}> 完成")
            else:
                # 如果没有找到处理器，处理其子元素
                if self.debug_mode:
                    self.logger.debug(f"未找到元素 <{element.name}> 的处理器，处理其子元素")
                for child in element.children:
                    if isinstance(child, Tag):
                        self._process_element(child)
        except Exception as e:
            self.logger.error(f"处理元素 <{element.name}> 时发生错误: {str(e)}")
            # 继续处理，不中断转换过程
    
    def cleanup(self):
        """
        /**
         * 清理临时资源
         */
        """
        self.logger.debug("开始清理临时资源")
        # 清理HTML元素处理器的临时资源
        self.elements_processor.cleanup()
        self.logger.info("清理临时资源完成")