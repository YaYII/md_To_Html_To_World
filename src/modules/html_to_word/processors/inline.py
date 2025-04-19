"""
内联元素处理器模块
处理HTML内联元素，如加粗、斜体、链接等
"""

from typing import Dict, Any, Optional, List
from bs4 import Tag
from docx.text.paragraph import Paragraph

from .base import BaseElementProcessor

class InlineElementProcessor(BaseElementProcessor):
    """
    /**
     * 内联元素处理器
     * 
     * 处理HTML内联元素，如加粗、斜体、链接等
     */
    """
    
    def can_process(self, element: Tag) -> bool:
        """
        /**
         * 检查是否可以处理指定元素
         * 
         * @param {Tag} element - HTML元素
         * @returns {bool} 是否可以处理该元素
         */
        """
        inline_tags = ['strong', 'b', 'em', 'i', 'u', 'code', 'a', 'span', 'br']
        return element.name in inline_tags
    
    def process(self, element: Tag, paragraph: Optional[Paragraph] = None) -> Optional[Paragraph]:
        """
        /**
         * 处理内联元素
         * 
         * @param {Tag} element - 内联HTML元素
         * @param {Optional[Paragraph]} paragraph - 段落对象，如果为None则创建新段落
         * @returns {Optional[Paragraph]} 处理后的段落对象
         */
        """
        if not paragraph:
            paragraph = self.document.add_paragraph()
            self.style_manager.apply_paragraph_format(paragraph)
        
        if element.name == 'br':
            # 处理换行
            paragraph.add_run('\n')
        elif element.name in ['strong', 'b']:
            # 处理粗体
            run = paragraph.add_run(element.get_text())
            run.bold = True
            self.style_manager.apply_default_style(run)
        elif element.name in ['em', 'i']:
            # 处理斜体
            run = paragraph.add_run(element.get_text())
            run.italic = True
            self.style_manager.apply_default_style(run)
        elif element.name == 'u':
            # 处理下划线
            run = paragraph.add_run(element.get_text())
            run.underline = True
            self.style_manager.apply_default_style(run)
        elif element.name == 'code':
            # 处理行内代码
            run = paragraph.add_run(element.get_text())
            self.style_manager.apply_code_style(run)
        elif element.name == 'a':
            # 处理链接
            url = element.get('href', '')
            text = element.get_text()
            if url and text:
                run = paragraph.add_run(text)
                self.style_manager.apply_link_style(run)
        elif element.name == 'span':
            # 处理span，根据样式属性应用不同格式
            run = paragraph.add_run(element.get_text())
            self.style_manager.apply_default_style(run)
            
            # 根据style属性应用特殊样式
            if 'style' in element.attrs:
                style = element['style']
                if 'font-weight:bold' in style or 'font-weight: bold' in style:
                    run.bold = True
                if 'font-style:italic' in style or 'font-style: italic' in style:
                    run.italic = True
        
        return paragraph
    
    def process_inline_elements(self, element: Tag, paragraph: Paragraph) -> Paragraph:
        """
        /**
         * 处理元素中的所有内联内容
         * 
         * @param {Tag} element - 包含内联内容的HTML元素
         * @param {Paragraph} paragraph - 段落对象
         * @returns {Paragraph} 处理后的段落对象
         */
        """
        # 如果元素为空，直接返回
        if not element or not element.contents:
            return paragraph
            
        # 处理所有子内容
        for content in element.contents:
            if isinstance(content, Tag):
                # 根据不同标签类型进行处理
                if content.name in ['br']:
                    # 添加换行
                    paragraph.add_run('\n')
                elif content.name in ['strong', 'b']:
                    # 处理粗体元素
                    if content.contents:
                        # 如果包含嵌套元素，递归处理
                        if any(isinstance(child, Tag) for child in content.contents):
                            # 创建一个新的run用于包装嵌套元素
                            for child in content.contents:
                                if isinstance(child, Tag):
                                    self.process_inline_elements(child, paragraph)
                                else:
                                    # 处理文本内容
                                    run = paragraph.add_run(str(child))
                                    run.bold = True
                                    self.style_manager.apply_default_style(run)
                        else:
                            # 只有文本内容，直接处理
                            run = paragraph.add_run(content.get_text())
                            run.bold = True
                            self.style_manager.apply_default_style(run)
                elif content.name in ['em', 'i']:
                    # 处理斜体元素
                    if content.contents:
                        # 如果包含嵌套元素，递归处理
                        if any(isinstance(child, Tag) for child in content.contents):
                            # 创建一个新的run用于包装嵌套元素
                            for child in content.contents:
                                if isinstance(child, Tag):
                                    self.process_inline_elements(child, paragraph)
                                else:
                                    # 处理文本内容
                                    run = paragraph.add_run(str(child))
                                    run.italic = True
                                    self.style_manager.apply_default_style(run)
                        else:
                            # 只有文本内容，直接处理
                            run = paragraph.add_run(content.get_text())
                            run.italic = True
                            self.style_manager.apply_default_style(run)
                elif content.name == 'code':
                    # 处理行内代码
                    run = paragraph.add_run(content.get_text())
                    self.style_manager.apply_code_style(run)
                elif content.name == 'a':
                    # 处理链接
                    url = content.get('href', '')
                    text = content.get_text()
                    if url and text:
                        run = paragraph.add_run(text)
                        self.style_manager.apply_link_style(run)
                elif content.name in ['p', 'div', 'span']:
                    # 递归处理块级元素中的内联内容
                    self.process_inline_elements(content, paragraph)
                else:
                    # 递归处理其他标签
                    self.process_inline_elements(content, paragraph)
            else:
                # 处理纯文本
                text = str(content)
                if text.strip():
                    run = paragraph.add_run(text)
                    self.style_manager.apply_default_style(run)
        
        return paragraph 