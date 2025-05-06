"""
列表处理器模块
处理HTML列表元素
"""

from typing import Dict, Any, Optional, List, Union
from bs4 import Tag
from docx.text.paragraph import Paragraph

from .base import BaseProcessor
from .inline import InlineProcessor

class ListProcessor(BaseProcessor):
    """
    /**
     * 列表处理器
     * 
     * 处理HTML列表元素，如ul、ol和li
     */
    """
    
    def can_process(self, element: Tag) -> bool:
        """
        /**
         * 检查是否可以处理指定元素
         * 
         * @param {Tag} element - HTML元素
         * @returns {bool} 是否可以处理该元素
         */
        """
        list_tags = ['ul', 'ol', 'li']
        return element.name in list_tags
    
    def process(self, element: Tag, level: int = 0) -> Union[List[Paragraph], None]:
        """
        /**
         * 处理列表元素
         * 
         * @param {Tag} element - 列表HTML元素
         * @param {int} level - 列表嵌套级别
         * @returns {Union[List[Paragraph], None]} 处理后的段落对象列表或None
         */
        """
        if element.name in ['ul', 'ol']:
            return self._process_list(element, level)
        elif element.name == 'li':
            return self._process_list_item(element, level)
        
        return None
    
    def _process_list(self, element: Tag, level: int = 0) -> List[Paragraph]:
        """
        /**
         * 处理列表容器元素
         * 
         * @param {Tag} element - 列表容器HTML元素
         * @param {int} level - 列表嵌套级别
         * @returns {List[Paragraph]} 处理后的段落对象列表
         */
        """
        paragraphs = []
        is_ordered = element.name == 'ol'
        item_index = 1
        
        # 处理列表项
        for item in element.find_all('li', recursive=False):
            # 处理有序列表项
            if is_ordered:
                # 创建列表项段落
                p = self.document.add_paragraph()
                self.style_manager.apply_list_item_format(p, level)
                
                # 添加有序列表标记
                p.add_run(f"{item_index}. ")
                item_index += 1
                
                # 处理列表项内容
                self._process_list_item_content(item, p, level)
                
                paragraphs.append(p)
            # 处理无序列表项
            else:
                # 创建列表项段落
                p = self.document.add_paragraph()
                self.style_manager.apply_list_item_format(p, level)
                
                # 添加无序列表标记
                p.add_run("• ")
                
                # 处理列表项内容
                self._process_list_item_content(item, p, level)
                
                paragraphs.append(p)
                
            # 处理嵌套列表
            nested_lists = item.find_all(['ul', 'ol'], recursive=False)
            for nested_list in nested_lists:
                nested_paragraphs = self._process_list(nested_list, level + 1)
                paragraphs.extend(nested_paragraphs)
        
        return paragraphs
    
    def _process_list_item(self, element: Tag, level: int = 0) -> List[Paragraph]:
        """
        /**
         * 处理列表项元素
         * 
         * @param {Tag} element - 列表项HTML元素
         * @param {int} level - 列表嵌套级别
         * @returns {List[Paragraph]} 处理后的段落对象列表
         */
        """
        paragraphs = []
        
        # 创建列表项段落
        p = self.document.add_paragraph()
        self.style_manager.apply_list_item_format(p, level)
        
        # 添加列表项标记（默认无序列表）
        p.add_run("• ")
        
        # 处理列表项内容
        self._process_list_item_content(element, p, level)
        
        paragraphs.append(p)
        
        # 处理嵌套列表
        nested_lists = element.find_all(['ul', 'ol'], recursive=False)
        for nested_list in nested_lists:
            nested_paragraphs = self._process_list(nested_list, level + 1)
            paragraphs.extend(nested_paragraphs)
        
        return paragraphs
    
    def _process_list_item_content(self, element: Tag, paragraph: Paragraph, level: int):
        """
        /**
         * 处理列表项内容
         * 
         * @param {Tag} element - 列表项HTML元素
         * @param {Paragraph} paragraph - 段落对象
         * @param {int} level - 列表嵌套级别
         */
        """
        # 检查列表项内是否包含复杂元素，如段落、div等
        complex_elements = element.find_all(['p', 'div', 'blockquote'], recursive=False)
        
        if complex_elements:
            # 处理第一个复杂元素的内容
            first_element = complex_elements[0]
            inline_processor = InlineProcessor(self.document, self.style_manager)
            inline_processor.process_inline_elements(first_element, paragraph)
            
            # 其他复杂元素需要单独处理
            if len(complex_elements) > 1:
                from ..element_factory import ElementProcessorFactory
                factory = ElementProcessorFactory(self.document, self.style_manager)
                
                for complex_element in complex_elements[1:]:
                    processor = factory.get_processor(complex_element)
                    if processor:
                        processor.process(complex_element)
        else:
            # 处理普通内容
            # 检查列表项内是否包含strong标签
            strong_tags = element.find_all(['strong', 'b'])
            if strong_tags:
                # 处理列表项中的复杂内容（包含格式化标签）
                self._process_list_item_with_format(element, paragraph)
            else:
                # 处理简单文本内容
                inline_processor = InlineProcessor(self.document, self.style_manager)
                inline_processor.process_inline_elements(element, paragraph)
    
    def _process_list_item_with_format(self, element: Tag, paragraph: Paragraph):
        """
        /**
         * 处理包含格式化标签的列表项
         * 
         * @param {Tag} element - 列表项元素
         * @param {Paragraph} paragraph - Word段落对象
         */
        """
        # 收集所有需要单独处理的子元素
        for content in element.children:
            if isinstance(content, Tag):
                if content.name in ['ul', 'ol']:
                    # 嵌套列表会在外部处理
                    continue
                elif content.name in ['strong', 'b']:
                    # 处理粗体文本
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                    self.style_manager.apply_default_style(run)
                elif content.name in ['em', 'i']:
                    # 处理斜体文本
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                    self.style_manager.apply_default_style(run)
                elif content.name == 'code':
                    # 处理行内代码
                    run = paragraph.add_run(content.get_text())
                    self.style_manager.apply_code_style(run)
                elif content.name == 'br':
                    # 处理换行
                    paragraph.add_run('\n')
                else:
                    # 使用内联处理器处理其他标签
                    inline_processor = InlineProcessor(self.document, self.style_manager)
                    inline_processor.process(content, paragraph)
            else:
                # 处理纯文本
                text = str(content)
                if text.strip():
                    run = paragraph.add_run(text)
                    self.style_manager.apply_default_style(run) 