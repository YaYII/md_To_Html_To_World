"""
表格处理器模块
处理HTML表格元素
"""

from typing import Dict, Any, Optional, List, Union, Tuple
from bs4 import Tag
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from .base import BaseElementProcessor
from .inline import InlineElementProcessor

# 定义边框样式常量
class BORDER_STYLE:
    """
    /**
     * 边框样式常量
     * 因为python-docx中没有WD_BORDER枚举，所以自定义一个
     */
    """
    NONE = "none"
    SINGLE = "single"
    THICK = "thick"
    DOUBLE = "double"
    DOTTED = "dotted"
    DASHED = "dashed"
    DOT_DASH = "dotDash"
    DOT_DOT_DASH = "dotDotDash"
    TRIPLE = "triple"

class TableProcessor(BaseElementProcessor):
    """
    /**
     * 表格处理器
     * 
     * 处理HTML表格元素，支持合并单元格和表格样式
     */
    """
    
    def can_process(self, element: Tag) -> bool:
        """
        /**
         * 检查是否可以处理指定元素
         * 
         * @param {Tag} element - HTML元素
         * @returns {bool} 是否可以处理该元素
         */
        """
        table_tags = ['table', 'tr', 'td', 'th', 'thead', 'tbody', 'tfoot']
        return element.name in table_tags
    
    def process(self, element: Tag) -> Union[Paragraph, None]:
        """
        /**
         * 处理表格元素
         * 
         * @param {Tag} element - 表格HTML元素
         * @returns {Union[Paragraph, None]} 处理后的段落对象或None
         */
        """
        # 只处理表格根元素
        if element.name != 'table':
            return None
            
        # 获取表格行和列信息
        rows = []
        
        # 1. 首先获取直接子元素中的行
        direct_rows = element.find_all('tr', recursive=False)
        if direct_rows:
            rows.extend(direct_rows)
        
        # 2. 获取thead中的行
        if element.thead:
            thead_rows = element.thead.find_all('tr', recursive=True)
            if thead_rows:
                rows.extend(thead_rows)
        
        # 3. 获取tbody中的行
        if element.tbody:
            tbody_rows = element.tbody.find_all('tr', recursive=True)
            if tbody_rows:
                rows.extend(tbody_rows)
                
        # 4. 获取tfoot中的行
        if element.tfoot:
            tfoot_rows = element.tfoot.find_all('tr', recursive=True)
            if tfoot_rows:
                rows.extend(tfoot_rows)
            
        if not rows:
            # 如果还没有找到行，尝试全局查找所有的行
            rows = element.find_all('tr')
            
        if not rows:
            return None
            
        # 计算列数
        max_cols = 0
        for row in rows:
            cols = row.find_all(['td', 'th'], recursive=False)
            max_cols = max(max_cols, len(cols))
            
        if max_cols == 0:
            return None
            
        # 创建表格
        table = self.document.add_table(rows=len(rows), cols=max_cols)
        
        # 全局设置所有单元格垂直居中 - 在应用样式前确保基础设置
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                cell = table.cell(i, j)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # 应用表格样式
        self._apply_table_style(table, element)
        
        # 填充表格内容
        self._fill_table_content(table, rows, max_cols)
        
        # 最后再次确保表头行垂直居中
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # 使用XML直接设置垂直对齐
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign_element = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                
                # 移除现有的垂直对齐设置（如果有）
                for old in tcPr.findall('.//{*}vAlign'):
                    tcPr.remove(old)
                
                # 添加新的垂直对齐设置
                tcPr.append(vAlign_element)
        
        # 在表格后添加空白段落
        spacer = self.document.add_paragraph()
        return spacer
    
    def _apply_table_style(self, table: Table, element: Tag):
        """
        /**
         * 应用表格样式
         * 
         * @param {Table} table - Word表格对象
         * @param {Tag} element - HTML表格元素
         */
        """
        # 获取增强表格样式配置
        enhanced_styles = self.style_manager.config.get('enhanced_table_styles', {})
        
        if self.debug_mode:
            self.logger.debug(f"应用表格样式，配置为: {enhanced_styles}")
        
        # 获取表格宽度百分比
        table_width_percent = 100  # 默认100%
        if 'width' in element.attrs:
            width_str = element['width']
            try:
                if width_str.endswith('%'):
                    table_width_percent = float(width_str.rstrip('%'))
                else:
                    # 如果是绝对宽度，默认使用100%
                    table_width_percent = 100
            except ValueError:
                table_width_percent = 100
        
        # 从配置中获取表格宽度
        table_width_percent = enhanced_styles.get('table_width_percent', table_width_percent)
        
        # 设置表格宽度 (相对于页面宽度的百分比)
        try:
            if table_width_percent:
                # 获取页面宽度
                section = self.document.sections[0]
                page_width = section.page_width
                margin_left = section.left_margin
                margin_right = section.right_margin
                available_width = page_width - margin_left - margin_right
                
                # 计算表格实际宽度
                table_width = Inches(float(available_width.inches) * table_width_percent / 100)
                
                # 设置表格宽度
                table.width = table_width
                if self.debug_mode:
                    self.logger.debug(f"设置表格宽度为页面可用宽度的 {table_width_percent}%")
        except Exception as e:
            if self.debug_mode:
                self.logger.error(f"设置表格宽度错误: {e}")
                
        # 应用边框样式
        try:
            # 获取边框样式配置
            border_style = enhanced_styles.get('border_style', 'single')
            border_size = enhanced_styles.get('border_size', 1)
            border_color = enhanced_styles.get('border_color', '000000')
            
            # 转换成python-docx能识别的边框样式字符串
            border_styles = {
                'none': BORDER_STYLE.NONE,
                'single': BORDER_STYLE.SINGLE, 
                'thick': BORDER_STYLE.THICK,
                'double': BORDER_STYLE.DOUBLE,
                'dotted': BORDER_STYLE.DOTTED,
                'dashed': BORDER_STYLE.DASHED,
                'dotDash': BORDER_STYLE.DOT_DASH,
                'dotDotDash': BORDER_STYLE.DOT_DOT_DASH,
                'triple': BORDER_STYLE.TRIPLE
            }
            
            # 默认使用单线边框
            doc_border_style = border_styles.get(border_style, BORDER_STYLE.SINGLE)
            
            # 解析颜色
            if border_color.startswith('#'):
                border_color = border_color[1:]
                
            # 为表格应用边框样式
            for row in table.rows:
                for cell in row.cells:
                    # 使用XML直接设置边框样式
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # 创建边框XML元素
                    borders_xml = f'<w:tcBorders {nsdecls("w")}>'
                    
                    # 添加四个方向的边框
                    for direction in ['top', 'left', 'bottom', 'right']:
                        borders_xml += f'<w:{direction} w:val="{doc_border_style}" w:sz="{border_size*4}" w:space="0" w:color="{border_color}"/>'
                    
                    borders_xml += '</w:tcBorders>'
                    
                    # 移除可能存在的旧边框设置
                    for old in tcPr.findall('.//{*}tcBorders'):
                        tcPr.remove(old)
                    
                    # 添加新的边框设置
                    tcPr.append(parse_xml(borders_xml))
            
            if self.debug_mode:
                self.logger.debug(f"已应用表格边框样式: {border_style}, 粗细: {border_size}pt, 颜色: #{border_color}")
        except Exception as e:
            if self.debug_mode:
                self.logger.error(f"设置表格边框错误: {e}")
        
        # 设置表格行高
        try:
            # 获取行高配置
            height_config = enhanced_styles.get('row_height', {})
            
            # 默认值
            default_row_height = height_config.get('default', 0.95)  # 默认行高，单位：厘米
            header_row_height = height_config.get('header', default_row_height)  # 表头行高，单位：厘米
            min_row_height = height_config.get('min', 0.5)  # 最小允许行高，单位：厘米
            max_row_height = height_config.get('max', 2.0)  # 最大允许行高，单位：厘米
            auto_adjust = height_config.get('auto_adjust', True)  # 默认启用自动调整行高
            
            # 如果配置是简单的数字，则使用该值作为所有行的高度
            if isinstance(enhanced_styles.get('cell_height'), (int, float)):
                default_row_height = float(enhanced_styles.get('cell_height'))
                header_row_height = default_row_height
            
            if self.debug_mode:
                self.logger.debug(f"行高配置: 默认={default_row_height}厘米, 表头={header_row_height}厘米, "
                                  f"最小={min_row_height}厘米, 最大={max_row_height}厘米, 自动调整={auto_adjust}")
            
            # 设置每一行的高度
            for i, row in enumerate(table.rows):
                # 确定当前行使用的行高值
                current_height = header_row_height if i == 0 else default_row_height
                
                # 确保高度在合理范围内
                current_height = max(min_row_height, min(current_height, max_row_height))
                
                # 设置行高
                row.height = Cm(current_height)
                
                # 使用XML直接设置高度类型为"确切值"(exact)而非"至少值"(atLeast)
                # 除非配置了自动调整行高，则使用"至少值"来允许根据内容自动扩展行高
                height_rule = "atLeast" if auto_adjust else "exact"
                
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(current_height*567.0)}" w:hRule="{height_rule}"/>')
                
                # 移除可能存在的旧设置
                for old in trPr.findall('.//{*}trHeight'):
                    trPr.remove(old)
                    
                # 添加新设置
                trPr.append(trHeight)
                
            if self.debug_mode:
                self.logger.debug(f"已设置表格行高，类型: {height_rule}")
        except Exception as e:
            if self.debug_mode:
                self.logger.error(f"设置表格行高错误: {str(e)}")
        
        # 设置背景颜色
        even_row_color = enhanced_styles.get('even_row_color')
        header_bg_color = enhanced_styles.get('header_bg_color')
        if self.debug_mode:
            self.logger.debug(f"表头背景色: {header_bg_color}, 偶数行背景色: {even_row_color}")
        
        # 应用表头背景色
        if header_bg_color and len(table.rows) > 0:
            try:
                self._set_row_bg_color(table, 0, header_bg_color)
                if self.debug_mode:
                    self.logger.debug("已应用表头背景色")
            except Exception as e:
                if self.debug_mode:
                    self.logger.error(f"设置表头背景色错误: {e}")
            
        # 应用隔行背景色
        if even_row_color:
            try:
                for i in range(1, len(table.rows), 2):
                    self._set_row_bg_color(table, i, even_row_color)
                if self.debug_mode:
                    self.logger.debug("已应用隔行背景色")
            except Exception as e:
                if self.debug_mode:
                    self.logger.error(f"设置隔行背景色错误: {e}")
        
        # 设置表头在分页时保持可见
        if enhanced_styles.get('keep_header_visible', True) and len(table.rows) > 0:
            try:
                # 获取表格第一行
                first_row = table.rows[0]._tr
                
                # 设置表格重复标题行属性
                tbl_pr = table._tbl.get_or_add_tblPr()
                
                # 添加表头行重复属性
                tbl_pr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                
                # 设置行属性不允许分页
                tr_pr = first_row.get_or_add_trPr()
                tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                if self.debug_mode:
                    self.logger.debug("已设置表头在分页时保持可见")
            except Exception as e:
                if self.debug_mode:
                    self.logger.error(f"设置表头分页属性错误: {e}")
            
        # 设置垂直对齐方式
        vertical_align = enhanced_styles.get('vertical_align', 'center')
        vertical_align_map = {
            'top': WD_CELL_VERTICAL_ALIGNMENT.TOP,
            'center': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            'bottom': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        }
        
        cell_v_align = vertical_align_map.get(vertical_align, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        try:
            # 设置所有单元格的垂直对齐方式
            for row in table.rows:
                for cell in row.cells:
                    # 使用python-docx API设置垂直对齐
                    cell.vertical_alignment = cell_v_align
                    
                    # 使用XML直接设置垂直对齐，确保所有Office版本兼容
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # 映射垂直对齐方式到XML值
                    xml_valign = {
                        WD_CELL_VERTICAL_ALIGNMENT.TOP: 'top',
                        WD_CELL_VERTICAL_ALIGNMENT.CENTER: 'center',
                        WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: 'bottom',
                        WD_CELL_VERTICAL_ALIGNMENT.BOTH: 'both'
                    }.get(cell_v_align, 'center')
                    
                    # 创建垂直对齐XML元素
                    vAlign_element = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{xml_valign}"/>')
                    
                    # 移除现有的垂直对齐设置（如果有）
                    for old in tcPr.findall('.//{*}vAlign'):
                        tcPr.remove(old)
                    
                    # 添加新的垂直对齐设置
                    tcPr.append(vAlign_element)
                    
            if self.debug_mode:
                self.logger.debug(f"已应用垂直对齐方式: {vertical_align} (XML: {xml_valign})")
        except Exception as e:
            if self.debug_mode:
                self.logger.error(f"设置垂直对齐方式错误: {e}")
                
        # 特别处理表头行，确保表头垂直居中
        try:
            if len(table.rows) > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    # 使用python-docx API设置垂直对齐
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    # 使用XML直接设置垂直对齐，加强处理确保最终生效
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    vAlign_element = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                    
                    # 移除现有的垂直对齐设置（如果有）
                    for old in tcPr.findall('.//{*}vAlign'):
                        tcPr.remove(old)
                    
                    # 添加新的垂直对齐设置
                    tcPr.append(vAlign_element)
                    
                if self.debug_mode:
                    self.logger.debug(f"已特别强化表头行的垂直居中设置")
        except Exception as e:
            if self.debug_mode:
                self.logger.error(f"强化设置表头行垂直对齐方式错误: {e}")
    
    def _set_row_bg_color(self, table: Table, row_index: int, color_str: str):
        """
        /**
         * 设置表格行的背景颜色
         * 
         * @param {Table} table - Word表格对象
         * @param {int} row_index - 行索引
         * @param {str} color_str - 颜色字符串，例如 '#RRGGBB'
         */
        """
        if row_index >= len(table.rows):
            if self.debug_mode:
                self.logger.debug(f"行索引超出范围: {row_index}, 表格总行数: {len(table.rows)}")
            return
            
        if self.debug_mode:
            self.logger.debug(f"设置行 {row_index} 背景色: {color_str}")
        
        try:
            # 解析颜色，支持 '#RRGGBB' 格式
            if color_str.startswith('#'):
                color_str = color_str[1:]
                
            # 解析16进制颜色值
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
            hex_color = f"{r:02x}{g:02x}{b:02x}"
            if self.debug_mode:
                self.logger.debug(f"解析的颜色: R={r}, G={g}, B={b}, HEX={hex_color}")
            
            # 获取行对象
            row = table.rows[row_index]
            
            # 为行中的每个单元格设置背景色
            for cell in row.cells:
                # 使用XML直接设置背景色，而不是通过python-docx API
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # 移除现有的单元格阴影设置（如果有）
                existing_shading = tcPr.find(qn('w:shd'))
                if existing_shading is not None:
                    tcPr.remove(existing_shading)
                
                # 创建并添加新的阴影元素 (使用w:fill属性设置背景色)
                shading_element = parse_xml(
                    f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
                )
                tcPr.append(shading_element)
            
            if self.debug_mode:
                self.logger.debug(f"已成功设置行 {row_index} 的背景色")
            
        except Exception as e:
            if self.debug_mode:
                self.logger.error(f"设置背景色错误: {e}")
            # 如果发生错误，使用默认浅灰色
            try:
                # 获取行对象
                row = table.rows[row_index]
                
                # 为行中的每个单元格设置背景色
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # 移除现有的单元格阴影设置（如果有）
                    existing_shading = tcPr.find(qn('w:shd'))
                    if existing_shading is not None:
                        tcPr.remove(existing_shading)
                    
                    # 使用默认的浅灰色
                    shading_element = parse_xml(
                        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="E7E6E6"/>'
                    )
                    tcPr.append(shading_element)
                
                if self.debug_mode:
                    self.logger.debug("已应用默认背景色")
            except Exception as ex:
                if self.debug_mode:
                    self.logger.error(f"应用默认背景色也失败: {ex}")
    
    def _fill_table_content(self, table: Table, rows: List[Tag], max_cols: int):
        """
        /**
         * 填充表格内容
         * 
         * @param {Table} table - Word表格对象
         * @param {List[Tag]} rows - 表格行元素列表
         * @param {int} max_cols - 最大列数
         */
        """
        # 从配置中获取行高
        enhanced_styles = self.style_manager.config.get('enhanced_table_styles', {})
        
        # 获取行高配置
        height_config = enhanced_styles.get('row_height', {})
        if isinstance(height_config, dict):
            default_row_height = height_config.get('default', 0.95)  # 默认行高，单位：厘米
            header_row_height = height_config.get('header', default_row_height)  # 表头行高，单位：厘米
            auto_adjust = height_config.get('auto_adjust', True)  # 默认启用自动调整行高
        else:
            # 兼容旧配置，直接使用cell_height作为默认行高
            default_row_height = enhanced_styles.get('cell_height', 0.95)  # 单位：厘米
            header_row_height = default_row_height
            auto_adjust = True  # 默认启用自动调整
        
        if self.debug_mode:
            self.logger.debug(f"填充表格内容，行数: {len(rows)}，列数: {max_cols}，"
                              f"默认行高: {default_row_height}厘米，表头行高: {header_row_height}厘米，"
                              f"自动调整: {auto_adjust}")
            
        # 创建单元格合并跟踪矩阵
        # 值为0表示正常可用单元格，值为1表示被合并的单元格
        merged_matrix = [[0 for _ in range(max_cols)] for _ in range(len(rows))]
        
        # 填充表格内容
        for i, row in enumerate(rows):
            # 确定当前行使用的行高值
            current_height = header_row_height if i == 0 else default_row_height
            
            # 设置行高为固定值或自动调整值
            try:
                table_row = table.rows[i]
                table_row.height = Cm(current_height)
                
                # 使用XML直接设置高度类型
                height_rule = "atLeast" if auto_adjust else "exact"
                
                tr = table_row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(current_height*567.0)}" w:hRule="{height_rule}"/>')
                
                # 移除可能存在的旧设置
                for old in trPr.findall('.//{*}trHeight'):
                    trPr.remove(old)
                    
                # 添加新设置
                trPr.append(trHeight)
                
                if self.debug_mode and i == 0:
                    self.logger.debug(f"设置第 {i+1} 行高度: {current_height}厘米, 类型: {height_rule}")
            except Exception as e:
                if self.debug_mode:
                    self.logger.error(f"在填充内容过程中设置行 {i} 高度时出错: {str(e)}")
            
            cells = row.find_all(['td', 'th'], recursive=False)
            is_header = row.parent.name == 'thead' or cells and cells[0].name == 'th'
            
            col_index = 0  # 实际列索引，会根据合并单元格情况调整
            for j, cell in enumerate(cells):
                # 跳过已被合并的单元格
                while col_index < max_cols and merged_matrix[i][col_index] == 1:
                    col_index += 1
                
                if col_index >= max_cols:
                    break
                    
                # 处理合并单元格
                rowspan = int(cell.get('rowspan', 1))
                colspan = int(cell.get('colspan', 1))
                
                if rowspan > 1 or colspan > 1:
                    # 标记被合并的单元格
                    for r in range(i, min(i + rowspan, len(rows))):
                        for c in range(col_index, min(col_index + colspan, max_cols)):
                            if r != i or c != col_index:
                                merged_matrix[r][c] = 1
                    
                    # 执行合并操作
                    end_row = min(i + rowspan - 1, len(rows) - 1)
                    end_col = min(col_index + colspan - 1, max_cols - 1)
                    self._merge_cells(table, i, col_index, end_row, end_col)
                
                # 获取单元格并设置内容
                table_cell = table.cell(i, col_index)
                # 确保每个单元格都垂直居中
                table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._process_cell_content(table_cell, cell, is_header)
                
                # 移动到下一列
                col_index += colspan
    
    def _merge_cells(self, table: Table, start_row: int, start_col: int, end_row: int, end_col: int):
        """
        /**
         * 合并表格单元格
         * 
         * @param {Table} table - Word表格对象
         * @param {int} start_row - 起始行索引
         * @param {int} start_col - 起始列索引
         * @param {int} end_row - 结束行索引
         * @param {int} end_col - 结束列索引
         */
        """
        # 确保索引在有效范围内
        row_count = len(table.rows)
        col_count = len(table.columns)
        
        if end_row >= row_count or end_col >= col_count:
            return
        
        # 执行单元格合并
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        try:
            start_cell.merge(end_cell)
        except Exception:
            # 如果合并失败，忽略错误继续执行
            pass
    
    def _process_cell_content(self, table_cell: _Cell, html_cell: Tag, is_header: bool):
        """
        /**
         * 处理单元格内容
         * 
         * @param {_Cell} table_cell - Word表格单元格对象
         * @param {Tag} html_cell - HTML表格单元格元素
         * @param {bool} is_header - 是否是表头单元格
         */
        """
        # 清空单元格默认内容
        table_cell.text = ""
        
        # 获取第一个段落
        paragraph = table_cell.paragraphs[0]
        
        # 从配置中获取表格样式
        enhanced_styles = self.style_manager.config.get('enhanced_table_styles', {})
        
        # 设置段落对齐方式，默认左对齐，仅表头默认居中
        default_align = enhanced_styles.get('text_align', 'left')
        alignment = 'center' if is_header else default_align
        
        # 从单元格style属性或class属性中提取对齐方式
        if 'style' in html_cell.attrs:
            style = html_cell['style']
            if 'text-align' in style:
                import re
                align_match = re.search(r'text-align:\s*(\w+)', style)
                if align_match:
                    alignment = align_match.group(1)
        
        # 从class中提取对齐方式
        if 'class' in html_cell.attrs:
            classes = html_cell['class']
            if isinstance(classes, list):
                for cls in classes:
                    if cls in ['left', 'center', 'right', 'justify']:
                        alignment = cls
                        break
            elif isinstance(classes, str):
                for cls in classes.split():
                    if cls in ['left', 'center', 'right', 'justify']:
                        alignment = cls
                        break
        
        # 应用段落格式
        self.style_manager.apply_paragraph_format(paragraph, alignment)
        
        # 设置单元格垂直对齐为居中 - 对所有单元格应用，不只是表头
        table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # 使用XML直接设置垂直对齐，确保所有Office版本兼容
        tc = table_cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign_element = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
        
        # 移除现有的垂直对齐设置（如果有）
        for old in tcPr.findall('.//{*}vAlign'):
            tcPr.remove(old)
        
        # 添加新的垂直对齐设置
        tcPr.append(vAlign_element)
        
        if self.debug_mode:
            self.logger.debug(f"处理单元格内容，是否表头: {is_header}, 对齐方式: {alignment}, 已设置垂直居中")
        
        # 检查单元格内容类型
        has_children = any(isinstance(child, Tag) for child in html_cell.children)
        
        if has_children:
            # 如果有子元素，使用内联处理器处理所有内容
            inline_processor = InlineElementProcessor(self.document, self.style_manager)
            inline_processor.process_inline_elements(html_cell, paragraph)
        else:
            # 简单文本内容处理
            run = paragraph.add_run(html_cell.get_text().strip())
            
            # 应用样式但保留表头设置
            if is_header:
                run.bold = True
                # 使用特殊的表头字体样式
                run.font.name = self.style_manager.heading_font
                run.font.size = Pt(self.style_manager.default_size)  # 表头使用默认字号
                # 设置中文字体
                if hasattr(run._element, 'rPr') and hasattr(run._element.rPr, 'rFonts'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style_manager.heading_font)
            else:
                # 非表头单元格使用默认样式
                self.style_manager.apply_default_style(run)
                
        # 确保表头有正确的样式
        if is_header:
            for run in paragraph.runs:
                # 保留原有设置，只修改必要的表头特性
                run.bold = True
                # 不直接应用heading_style，因为它可能会覆盖其他样式
                if not hasattr(run.font, 'name') or not run.font.name:
                    run.font.name = self.style_manager.heading_font
                    
                # 确保表头文字居中垂直对齐
                cell = paragraph._parent
                if hasattr(cell, 'vertical_alignment'):
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    # 使用XML直接设置垂直对齐，确保所有Office版本兼容
                    if hasattr(cell, '_tc'):
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        vAlign_element = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                        
                        # 移除现有的垂直对齐设置（如果有）
                        for old in tcPr.findall('.//{*}vAlign'):
                            tcPr.remove(old)
                        
                        # 添加新的垂直对齐设置
                        tcPr.append(vAlign_element) 