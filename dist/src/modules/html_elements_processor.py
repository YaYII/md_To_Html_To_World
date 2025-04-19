"""
HTML元素处理模块
提供处理HTML特殊元素的功能，如图片、代码高亮等
"""

import os
import re
import requests
from typing import Dict, Any, Optional, List, Tuple
from bs4 import Tag
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

class HtmlElementsProcessor:
    """
    /**
     * HTML特殊元素处理器
     * 
     * 处理HTML中的特殊元素，如图片、代码高亮等
     * 提供独立的处理方法，可被HTML到Word转换器调用
     */
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        /**
         * 初始化HTML元素处理器
         * 
         * @param {Dict[str, Any]} config - 配置参数字典
         */
        """
        self.config = config
        self.image_config = self.config.get('images', {})
        self.default_image_width = self.image_config.get('default_width', 10)  # 单位：厘米
        self.default_image_height = self.image_config.get('default_height', None)  # 自动计算高度
        self.max_image_width = self.image_config.get('max_width', 15)  # 单位：厘米
        self.temp_dir = None
    
    def process_image(self, img_tag: Tag, document: Document) -> Optional[Paragraph]:
        """
        /**
         * 处理图片元素，将图片插入Word文档
         * 
         * @param {Tag} img_tag - 图片标签元素
         * @param {Document} document - Word文档对象
         * @returns {Optional[Paragraph]} 包含图片的段落对象或None
         */
        """
        try:
            # 获取图片源
            src = img_tag.get('src', '')
            if not src:
                return None
                
            # 获取图片alt文本
            alt = img_tag.get('alt', '')
            
            # 创建临时目录（如果尚未创建）
            if self.temp_dir is None:
                self.temp_dir = tempfile.mkdtemp()
                
            # 下载或读取图片文件
            img_path = self._get_image_path(src)
            if not img_path:
                return None
                
            # 获取图片尺寸设置
            width_attr = img_tag.get('width', '')
            width = self._parse_size_attribute(width_attr, self.default_image_width)
            
            # 检查图片宽度是否超过最大值
            if width > self.max_image_width:
                width = self.max_image_width
                
            # 创建段落并插入图片
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 插入图片
            run = p.add_run()
            run.add_picture(img_path, width=Cm(width))
            
            # 如果有alt文本，添加图片说明
            if alt and self.image_config.get('show_caption', True):
                caption_p = document.add_paragraph(f"图 {alt}")
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.paragraph_format.space_after = Pt(12)
                
            return p
            
        except Exception as e:
            # 图片处理失败时返回None
            print(f"图片处理失败: {str(e)}")
            return None
    
    def _get_image_path(self, src: str) -> Optional[str]:
        """
        /**
         * 获取图片路径，支持本地文件和远程URL
         * 
         * @param {str} src - 图片源路径
         * @returns {Optional[str]} 图片的本地路径或None
         */
        """
        # 检查是否为本地文件路径
        if os.path.exists(src):
            return src
            
        # 检查是否为URL
        if src.startswith(('http://', 'https://')):
            try:
                # 从URL下载图片
                response = requests.get(src, stream=True, timeout=10)
                if response.status_code == 200:
                    # 提取文件名
                    filename = os.path.basename(src)
                    if not filename:
                        filename = f"image_{hash(src)}.jpg"
                        
                    # 保存到临时文件
                    img_path = os.path.join(self.temp_dir, filename)
                    with open(img_path, 'wb') as f:
                        for chunk in response.iter_content(1024):
                            f.write(chunk)
                    return img_path
            except Exception:
                pass
                
        # 如果是相对路径，尝试在指定目录中查找
        image_dirs = self.image_config.get('search_dirs', [])
        for dir_path in image_dirs:
            full_path = os.path.join(dir_path, src)
            if os.path.exists(full_path):
                return full_path
                
        return None
    
    def _parse_size_attribute(self, size_attr: str, default_value: float) -> float:
        """
        /**
         * 解析HTML中的尺寸属性
         * 
         * @param {str} size_attr - 尺寸属性字符串，可能包含单位
         * @param {float} default_value - 默认值
         * @returns {float} 解析后的尺寸值（厘米）
         */
        """
        if not size_attr:
            return default_value
            
        # 尝试解析纯数字
        try:
            return float(size_attr) / 40  # 假设为像素，转换为厘米（近似值）
        except ValueError:
            pass
            
        # 解析带单位的值
        match = re.match(r'(\d+\.?\d*)(\w*)', size_attr)
        if match:
            value, unit = match.groups()
            try:
                value = float(value)
                # 根据单位进行转换
                if unit.lower() == 'px':
                    return value / 40  # 像素转厘米
                elif unit.lower() == 'cm':
                    return value
                elif unit.lower() == 'in':
                    return value * 2.54  # 英寸转厘米
                elif unit.lower() == 'pt':
                    return value / 28.3  # 点转厘米
                elif unit.lower() == '%':
                    # 百分比值，以最大宽度为基准
                    return (value / 100) * self.max_image_width
            except ValueError:
                pass
                
        return default_value
    
    def process_code_highlight(self, code_tag: Tag, language: Optional[str] = None) -> str:
        """
        /**
         * 处理代码高亮
         * 
         * @param {Tag} code_tag - 代码标签元素
         * @param {Optional[str]} language - 代码语言
         * @returns {str} 格式化后的代码文本
         */
        """
        # 获取代码文本
        code_text = code_tag.get_text()
        
        # 移除首尾空行
        code_text = code_text.strip('\n')
        
        # 检测语言（如果没有提供）
        if not language and 'class' in code_tag.attrs:
            class_list = code_tag['class']
            for cls in class_list:
                if cls.startswith('language-'):
                    language = cls[9:]
                    break
        
        # 进行格式化（简单实现，可以扩展为更复杂的语法高亮）
        formatted_code = code_text
        
        return formatted_code
    
    def process_math_formula(self, math_tag: Tag) -> str:
        """
        /**
         * 处理数学公式
         * 
         * @param {Tag} math_tag - 数学公式标签元素
         * @returns {str} 格式化后的公式文本
         */
        """
        # 获取公式文本
        formula_text = math_tag.get_text()
        
        # 简单处理，只返回原始公式文本
        # 实际应用中可以使用更复杂的公式转换方法
        return f"[公式: {formula_text}]"
    
    def cleanup(self):
        """
        /**
         * 清理临时资源
         */
        """
        # 删除临时目录
        if self.temp_dir and os.path.exists(self.temp_dir):
            import shutil
            try:
                shutil.rmtree(self.temp_dir)
            except Exception:
                pass 